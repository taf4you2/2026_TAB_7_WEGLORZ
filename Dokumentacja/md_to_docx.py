#!/usr/bin/env python3
"""Prosty konwerter RAPORT.md -> FULL_DOCUMENTATION.docx.
Obsługuje: nagłówki (#..####), tabele GFM, listy (-, 1.), cytaty (>),
bloki kodu ```...```, pogrubienia **..**, kod inline `..`, linie poziome (---).
"""
import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = sys.argv[1] if len(sys.argv) > 1 else "Dokumentacja/RAPORT.md"
DST = sys.argv[2] if len(sys.argv) > 2 else "FULL_DOCUMENTATION.docx"

doc = Document()

# Bazowy font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")

def add_inline(paragraph, text):
    """Dodaje tekst z obsługą **bold** i `code` jako osobne runy."""
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.color.rgb = RGBColor(0xB0, 0x30, 0x60)
        else:
            paragraph.add_run(part)

def strip_md(text):
    """Usuwa zapis linków [tekst](url) -> tekst, na potrzeby komórek tabeli."""
    return re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)

with open(SRC, encoding="utf-8") as f:
    lines = f.read().split("\n")

i = 0
n = len(lines)
while i < n:
    line = lines[i]

    # Blok kodu
    if line.startswith("```"):
        i += 1
        code = []
        while i < n and not lines[i].startswith("```"):
            code.append(lines[i])
            i += 1
        i += 1  # zamykające ```
        p = doc.add_paragraph()
        r = p.add_run("\n".join(code))
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        # delikatne tło przez shading byłoby zbyt rozwlekłe; zostawiamy monospace
        continue

    # Tabela GFM
    if "|" in line and i + 1 < n and re.match(r"^\s*\|?[\s:|-]+\|?\s*$", lines[i + 1]) and "-" in lines[i + 1]:
        header = [c.strip() for c in line.strip().strip("|").split("|")]
        i += 2  # nagłówek + separator
        rows = []
        while i < n and "|" in lines[i] and lines[i].strip():
            rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
            i += 1
        table = doc.add_table(rows=1, cols=len(header))
        table.style = "Light Grid Accent 1"
        for j, h in enumerate(header):
            cell = table.rows[0].cells[j]
            cell.paragraphs[0].text = ""
            add_inline(cell.paragraphs[0], strip_md(h))
            for run in cell.paragraphs[0].runs:
                run.bold = True
        for row in rows:
            cells = table.add_row().cells
            for j in range(len(header)):
                val = row[j] if j < len(row) else ""
                cells[j].paragraphs[0].text = ""
                add_inline(cells[j].paragraphs[0], strip_md(val))
        doc.add_paragraph()
        continue

    # Linia pozioma
    if re.match(r"^\s*---+\s*$", line):
        i += 1
        continue

    # Nagłówki
    m = re.match(r"^(#{1,6})\s+(.*)$", line)
    if m:
        level = len(m.group(1))
        doc.add_heading(strip_md(m.group(2)), level=min(level, 4))
        i += 1
        continue

    # Cytat
    if line.startswith(">"):
        p = doc.add_paragraph(style="Intense Quote")
        add_inline(p, line.lstrip(">").strip())
        i += 1
        continue

    # Lista punktowana
    m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
    if m:
        p = doc.add_paragraph(style="List Bullet")
        add_inline(p, m.group(2))
        i += 1
        continue

    # Lista numerowana
    m = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
    if m:
        p = doc.add_paragraph(style="List Number")
        add_inline(p, m.group(2))
        i += 1
        continue

    # Pusta linia
    if not line.strip():
        i += 1
        continue

    # Zwykły akapit
    p = doc.add_paragraph()
    add_inline(p, line)
    i += 1

doc.save(DST)
print(f"Zapisano: {DST}")
